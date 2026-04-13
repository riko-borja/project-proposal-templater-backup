from flask import Flask, request, jsonify, send_file
from docx import Document
import requests
import io
import logging
import re
from flask_cors import CORS
import os

app = Flask(__name__)
CORS(app)  # Enable CORS for all routes

# Configure logging to include timestamps and log levels
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s [%(levelname)s] %(message)s',
    handlers=[
        logging.StreamHandler()
    ]
)

# Configuration for Node.js backend
NODE_BACKEND_URL = os.getenv('NODE_BACKEND_URL', 'http://127.0.0.1:5001').rstrip('/')

def fetch_document_from_node_backend(fileId, bucketName, userId, isModified=False):
    """
    Fetch the document from the Node.js backend.
    If `isModified` is True, fetch the modified document; otherwise, fetch the original.
    """
    if isModified:
        # New endpoint for fetching the modified document
        node_backend_url = f"{NODE_BACKEND_URL}/api/files/fetch-modified-template"
        params = {
            "userId": userId,
            "originalFileId": fileId  # Using `originalFileId` to link back to the original document
        }
    else:
        # Existing endpoint for fetching the original document
        node_backend_url = f"{NODE_BACKEND_URL}/api/files/fetch-template-content"
        params = {
            "fileId": fileId,
            "bucketName": bucketName
        }

    try:
        # Log the request details for debugging
        logging.debug(f"Fetching document from Node.js: {node_backend_url} with params {params}")

        response = requests.get(node_backend_url, params=params)
        response.raise_for_status()  # Raise an HTTPError for bad responses (4xx and 5xx)

        # Wrap the bytes content in a BytesIO object
        file_stream = io.BytesIO(response.content)
        file_stream.seek(0)  # Ensure the stream is at the beginning
        logging.debug("Successfully fetched and wrapped the document in BytesIO.")
        return file_stream
    except requests.exceptions.HTTPError as e:
        logging.error(f"HTTPError while fetching document: {e}")
        raise e
    except requests.exceptions.RequestException as e:
        logging.error(f"RequestException while fetching document: {e}")
        raise e


def extract_cell_text(cell):
    """
    Extracts and concatenates text from all paragraphs in a cell.
    """
    paragraphs = cell.paragraphs
    texts = [paragraph.text for paragraph in paragraphs]
    return '\n'.join(texts).strip()

def get_full_text(paragraph):
    """
    Reconstructs the full text of a paragraph by concatenating all runs.
    """
    return ''.join(run.text for run in paragraph.runs).strip()

def parse_table(table):
    """
    Parses a single table and extracts rows and cells.
    Identifies placeholders in the format {placeholder_id}.
    Handles colspan, rowspan, and nested tables.
    """
    table_content = []
    rowspan_map = {}  # Keeps track of row spans for cells

    for row_idx, row in enumerate(table.rows):
        row_content = []
        col_idx = 0  # Initialize column index
        while col_idx < len(row.cells):
            cell = row.cells[col_idx]

            # Adjust column index based on rowspan_map
            while (row_idx, col_idx) in rowspan_map:
                cell_placeholder = rowspan_map.pop((row_idx, col_idx))
                row_content.append(cell_placeholder)
                col_idx += 1

            cell_text = extract_cell_text(cell)
            logging.debug(f"Processing cell [{row_idx}][{col_idx}]: '{cell_text}'")

            # Detect placeholders
            placeholder_match = re.search(r'\{(.*?)\}', cell_text)
            is_placeholder = bool(placeholder_match)
            placeholder_id = placeholder_match.group(1) if is_placeholder else None

            # Detect colspan and rowspan
            colspan = 1  # Default colspan
            rowspan = 1  # Default rowspan

            # Handle colspan
            grid_span = cell._tc.tcPr.gridSpan
            if grid_span is not None:
                try:
                    colspan = int(grid_span.val)
                    logging.debug(f"Cell [{row_idx}][{col_idx}] has colspan: {colspan}")
                except (ValueError, TypeError):
                    logging.warning(f"Invalid gridSpan value in cell: {cell_text}")
                    colspan = 1

            # Handle rowspan
            v_merge = cell._tc.tcPr.vMerge
            if v_merge is not None:
                if v_merge.val == 'restart':
                    # Determine rowspan by checking subsequent rows
                    rowspan = 1
                    for i in range(1, 100):  # Arbitrary limit to prevent infinite loops
                        try:
                            next_row = table.rows[row_idx + i]
                            next_cell = next_row.cells[col_idx]
                            next_v_merge = next_cell._tc.tcPr.vMerge
                            if next_v_merge is not None and next_v_merge.val == 'continue':
                                rowspan += 1
                                logging.debug(f"Cell [{row_idx}][{col_idx}] spans to row {row_idx + i}")
                            else:
                                break
                        except IndexError:
                            break

                    if rowspan > 1:
                        for i in range(1, rowspan):
                            rowspan_map[(row_idx + i, col_idx)] = None  # Placeholder for merged cells
                        logging.debug(f"Cell [{row_idx}][{col_idx}] has rowspan: {rowspan}")
                else:
                    # Continuation of a rowspan; already handled
                    logging.debug(f"Cell [{row_idx}][{col_idx}] is a continuation of rowspan.")
                    col_idx += 1
                    continue

            # Check for nested tables within the cell
            nested_tables = cell.tables
            if nested_tables:
                logging.debug(f"Found nested table in cell [{row_idx}][{col_idx}]. Processing nested table.")
                nested_table_content = []
                for nested_table in nested_tables:
                    parsed_nested_table = parse_table(nested_table)
                    nested_table_content.append(parsed_nested_table)
                cell_data = {
                    'content': cell_text,
                    'colspan': colspan,
                    'rowspan': rowspan,
                    'isPlaceholder': is_placeholder,
                    'id': placeholder_id,
                    'value': '',
                    'nestedTables': nested_table_content  # Include nested tables if any
                }
            else:
                cell_data = {
                    'content': cell_text,
                    'colspan': colspan,
                    'rowspan': rowspan,
                    'isPlaceholder': is_placeholder,
                    'id': placeholder_id,
                    'value': ''  # Placeholder for user input
                }

            row_content.append(cell_data)

            # Handle colspan by adding placeholders
            for _ in range(1, colspan):
                row_content.append(None)  # Placeholder for merged columns

            col_idx += 1  # Move to the next cell

        # Append the processed row to table_content
        table_content.append(row_content)
        logging.debug(f"Row {row_idx} content: {row_content}")

    return table_content

def parse_document(file_stream):
    """
    Parses the .docx file to extract tables wrapped by {TABLE_START: TAGNAME} and {TABLE_END: TAGNAME}.
    Handles nested tables within other tables.
    """
    try:
        document = Document(file_stream)
        logging.debug("Opened the .docx document successfully.")
    except Exception as e:
        logging.exception(f"Failed to open the document: {e}")
        raise e

    tables_data = []
    table_name = None
    collecting = False
    current_table_data = None




    def get_paragraph_text(element):
        """
        Extracts text from a w:p element.
        """
        text = ''
        for child in element.iter():
            if child.tag.endswith('t'):  # Text
                text += child.text or ''
            elif child.tag.endswith('tab'):  # Tab
                text += '\t'
            elif child.tag.endswith('br'):  # Line Break
                text += '\n'
        return text.strip()


    def parse_nested_elements(elements):
        """Recursively parse elements, including nested tables."""
        nonlocal collecting, current_table_data, table_name

        for element in elements:
            if element.tag.endswith('p'):
                para_text = get_paragraph_text(element)
                logging.debug(f"Processing paragraph: '{para_text}'")
                # Detect start tag
                start_tag_match = re.match(r'\{TABLE_START:\s*(.+?)\s*\}', para_text, re.IGNORECASE)
                if start_tag_match:
                    table_name = start_tag_match.group(1).strip()
                    collecting = True
                    current_table_data = {
                        'name': table_name,
                        'rows': []
                    }
                    logging.debug(f"Started collecting table: {table_name}")
                    continue
                # Detect end tag
                end_tag_match = re.match(r'\{TABLE_END:\s*(.+?)\s*\}', para_text, re.IGNORECASE)
                if end_tag_match:
                    end_table_name = end_tag_match.group(1).strip()
                    if end_table_name == table_name:
                        tables_data.append(current_table_data)
                        logging.debug(f"Finished collecting table: {table_name}")
                        collecting = False
                        table_name = None
                        current_table_data = None
                    else:
                        logging.warning(f"End table name '{end_table_name}' does not match current table name '{table_name}'")
                    continue

            elif element.tag.endswith('tbl'):
                tbl = element
                if collecting and table_name:
                    logging.debug(f"Found table element for: {table_name}")
                    try:
                        temp_doc = Document()
                        temp_doc._body._body.append(tbl)
                        parsed_table = parse_table(temp_doc.tables[0])
                        current_table_data['rows'].extend(parsed_table)
                        logging.debug(f"Collected table data for: {table_name}")
                    except Exception as e:
                        logging.exception(f"Failed to parse table '{table_name}': {e}")
                else:
                    logging.debug("Found a table outside of any collecting context. Skipping processing.")
                # Recursively process elements inside the table
                parse_nested_elements(tbl)
            else:
                # Recursively process any child elements
                parse_nested_elements(element)


    # Start parsing from the main document body
    parse_nested_elements(document.element.body)

    return tables_data

def traverse_elements(element):
    """
    Recursively traverses all child elements, yielding each one.
    """
    for child in element:
        yield child
        yield from traverse_elements(child)



def _element_to_paragraph(element, document):
    """
    Converts an element to a paragraph object if possible.
    """
    for para in document.paragraphs:
        if para._p is element:
            return para
    return None


@app.route('/api/parse-document', methods=['GET'])
def parse_document_route():
    """
    Endpoint to parse a .docx document and extract table data and tags.
    """
    file_id = request.args.get('fileId')
    bucket_name = request.args.get('bucketName')
    user_id = request.args.get('userId')  # Fetch userId from request

    logging.debug(f"Received request in Flask: fileId={file_id}, bucketName={bucket_name}")

    if not file_id or not bucket_name:
        logging.error("Missing fileId or bucketName parameter.")
        return jsonify({'error': 'Missing fileId or bucketName parameter'}), 400

    try:
        # Fetch the document from Node.js backend
        file_stream = fetch_document_from_node_backend(file_id, bucket_name, user_id)

        # Parse the document for tables
        tables_data = parse_document(file_stream)

        # Parse the document for tags, tooltips, and hyperlinks
        file_stream.seek(0)  # Reset stream to the beginning
        collated_tags = parse_document_tags(file_stream)

        logging.debug(f"Parsed tables: {tables_data}")
        logging.debug(f"Collated tags: {collated_tags}")

        # Include both tables and tags in the response
        return jsonify({'tables': tables_data, 'tags': collated_tags}), 200

    except requests.exceptions.HTTPError as e:
        logging.error(f"HTTPError while fetching document: {e}")
        return jsonify({'error': 'Failed to fetch the document from Node.js backend.', 'details': str(e)}), 500
    except Exception as e:
        logging.exception("Error parsing the document")
        return jsonify({'error': 'Failed to parse the document.', 'details': str(e)}), 500



def parse_document_tags(file_stream):
    """
    Parses document text for table tags, hyperlinks, and tooltips from the entire document,
    including tables, paragraphs, and other elements.
    """
    document = Document(file_stream)
    collated_tags = {}

    # Updated tag patterns
    tag_patterns = {
        'table_hyperlink': r'\{table_hyperlink:([^|]+)\|([^\}]+)\}',
        'table_tooltip': r'\{table_tooltip:([^|]+)\|([^\}]+)\}',
        'table_hyperlink_tooltip': r'\{table_hyperlink_tooltip:([^|]+)\|([^\}]+)\}',
        'table_start': r'\{TABLE_START:([^\}]+)\}',
        'table_end': r'\{TABLE_END:([^\}]+)\}'
    }

    # Updated process_text_for_tags function
    def process_text_for_tags(text):
        """
        Extracts specific tags with required structures and updates collated_tags.
        """
        for tag_type, pattern in tag_patterns.items():
            for match in re.finditer(pattern, text):
                if tag_type in ['table_start', 'table_end']:
                    tag_name = match.group(1).strip()
                    # Store the tag with an empty value
                    collated_tags[f'{tag_type.upper()}:{tag_name}'] = ""
                elif tag_type in ['table_hyperlink', 'table_tooltip', 'table_hyperlink_tooltip']:
                    tag_name = match.group(1).strip()
                    tag_value = match.group(2).strip()
                    # Store the tag with the actual value
                    collated_tags[f'{tag_type}:{tag_name}'] = tag_value
                    # Store the tag with the value in the key and an empty value
                    collated_tags[f'{tag_type}:{tag_name}|{tag_value}'] = ""

    # Process all paragraphs in the document
    for para in document.paragraphs:
        process_text_for_tags(para.text)

    # Process all tables in the document
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                process_text_for_tags(cell.text)

    logging.debug(f"Collated tags: {collated_tags}")
    return collated_tags




def get_full_text_from_stream(file_stream):
    """
    Extracts the full text from the .docx file stream.
    """
    try:
        document = Document(file_stream)
        paragraphs = [para.text for para in document.paragraphs if para.text.strip()]
        return "\n".join(paragraphs)
    except Exception as e:
        logging.exception(f"Failed to extract full text: {e}")
        raise e


# Directory to save processed files
PROCESSED_DOCS_DIR = "processed_docs"
os.makedirs(PROCESSED_DOCS_DIR, exist_ok=True)

# Function to remove empty rows from tables (from your existing Python app)
def remove_empty_rows(doc):
    """
    Removes empty rows from all tables (including nested ones) in the given Word document.
    """
    for table in doc.tables:
        process_table(table)
        for row in table.rows:
            for cell in row.cells:
                for nested_table in cell.tables:
                    process_table(nested_table)

def process_table(table):
    """
    Processes a single table and removes rows that are completely empty.
    """
    rows_to_remove = []
    for row in table.rows:
        if all(is_cell_empty(cell) for cell in row.cells):
            rows_to_remove.append(row)
    for row in rows_to_remove:
        table._element.remove(row._element)

def is_cell_empty(cell):
    """
    Checks if a table cell is empty by removing all whitespace and non-visible characters.
    """
    text = cell.text.strip()
    text = re.sub(r'\s+', '', text)  # Remove all kinds of whitespace
    return len(text) == 0

@app.route('/api/remove-empty-rows', methods=['POST'])
def remove_empty_rows_route():
    # Log the request files to ensure 'file' is being sent
    logging.debug("Request files keys: %s", list(request.files.keys()))
    """
    Receives a document, processes it to remove empty rows from tables, 
    and returns the modified document to the frontend.
    """
    if 'file' not in request.files:
        logging.error("No 'file' key in the request.")
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if not file:
        return jsonify({'error': 'No file provided'}), 400

    # Load the document
    try:
        doc = Document(io.BytesIO(file.read()))  # Read the document into memory
        logging.debug("Document loaded successfully.")
    except Exception as e:
        logging.error(f"Error loading the document: {e}")
        return jsonify({'error': 'Failed to load the document'}), 500
    
    # Process the document to remove empty rows
    remove_empty_rows(doc)

    # Save the modified document to a BytesIO object
    modified_doc_stream = io.BytesIO()
    doc.save(modified_doc_stream)
    modified_doc_stream.seek(0)  # Rewind the stream to the beginning

    # Log the size of the document to ensure it's a valid file
    logging.debug(f"Sending the modified document, size: {len(modified_doc_stream.getvalue())} bytes")

    
    # Send the modified document back as a response
    return send_file(modified_doc_stream, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name='Modified_GeneratedDocument_SOW.docx')


if __name__ == '__main__':
    app.run(host='127.0.0.1', port=5000, debug=True)
